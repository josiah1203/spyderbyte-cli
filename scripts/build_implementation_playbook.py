from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "AGENTIC_PLATFORM_IMPLEMENTATION_PLAYBOOK.md"
OUTPUT = ROOT / "Agentic_Platform_Implementation_Playbook.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "9B761A"
RED = "9B1C1C"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color="C7CDD5", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def apply_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_inline_markdown(paragraph, text: str, default_size=11, default_color=None) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|https?://\S+)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=max(default_size - 1, 8), color=DARK_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=default_size, color=default_color, italic=True)
        else:
            run = paragraph.add_run(token.rstrip(".,;"))
            set_font(run, size=default_size, color=BLUE)
            run.underline = True
            if len(token.rstrip(".,;")) != len(token):
                tail = paragraph.add_run(token[-1])
                set_font(tail, size=default_size, color=default_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=default_size, color=default_color)


def paragraph_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def set_paragraph_border_bottom(paragraph, color=BLUE, size="10", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_callout(doc, text: str, label: str = "IMPLEMENTATION AUTHORITY") -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="B9C7D8", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_font(r, size=9, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_inline_markdown(p2, text, default_size=10.5, default_color=INK)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    configs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (11, INK, 8, 4),
    }
    for name, (size, color, before, after) in configs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_definition(abs_id: int, num_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), fmt)
            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text if fmt == "bullet" else f"%{level + 1}.")
            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "270")
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), "80")
            spacing.set(qn("w:line"), "300")
            spacing.set(qn("w:lineRule"), "auto")
            p_pr.extend([tabs, ind, spacing])
            lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
            if fmt == "bullet":
                r_pr = OxmlElement("w:rPr")
                fonts = OxmlElement("w:rFonts")
                fonts.set(qn("w:ascii"), "Symbol")
                fonts.set(qn("w:hAnsi"), "Symbol")
                r_pr.append(fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def new_numbering_instance(doc: Document, base_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    base = None
    for node in numbering.findall(qn("w:num")):
        if int(node.get(qn("w:numId"))) == base_num_id:
            base = node
            break
    if base is None:
        raise ValueError(f"Missing numbering definition {base_num_id}")
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_id = max(existing or [0]) + 1
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), abstract_id)
    num.append(abs_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return new_id


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    r = kicker.add_run("IMPLEMENTATION PLAYBOOK")
    set_font(r, size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Agentic ML/Data Platform")
    set_font(r, size=30, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    r = subtitle.add_run("End-to-End Runtime and Harness Implementation")
    set_font(r, size=16, color=DARK_BLUE)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(30)
    set_paragraph_border_bottom(rule, color=GOLD, size="12", space="5")

    meta = doc.add_table(rows=4, cols=2)
    apply_table_geometry(meta, [2160, 7200], indent_dxa=120)
    set_table_borders(meta, color=WHITE, size="0")
    rows = [
        ("STATUS", "Baseline implementation authority"),
        ("VERSION", "1.0"),
        ("PREPARED", "August 2, 2026"),
        ("REPOSITORY", "Greenfield baseline"),
    ]
    for idx, (label, value) in enumerate(rows):
        p1 = meta.cell(idx, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(label)
        set_font(r1, size=9, color=MUTED, bold=True)
        p2 = meta.cell(idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(value)
        set_font(r2, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("For implementation by Codex and the platform engineering team")
    set_font(r, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph("Document Map", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph("Use this map to locate the active work phase, contracts, gates, and verification requirements.")
    intro.paragraph_format.space_after = Pt(12)
    for text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_font(r, size=10.5, color=DARK_BLUE, bold=text.startswith("Appendix"))
    doc.add_page_break()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D3DAE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line or " ")
        set_font(r, name="Consolas", size=8.5, color="263442")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    if cols == 2:
        widths = [2400, 6960]
    elif cols == 3:
        widths = [2600, 1800, 4960]
    elif cols == 4:
        widths = [1500, 3300, 1400, 3160]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    apply_table_geometry(table, widths, indent_dxa=120)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        prevent_row_split(table.rows[r_idx])
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            value = row[c_idx] if c_idx < len(row) else ""
            add_inline_markdown(p, value, default_size=9.2, default_color=INK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(DARK_BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def render_markdown(doc: Document, text: str, bullet_num_id: int, decimal_num_id: int) -> None:
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    skipped_frontmatter = True
    previous_list_kind = None
    active_decimal_id = decimal_num_id
    while i < len(lines):
        line = lines[i].rstrip()
        if skipped_frontmatter:
            if line.startswith("## 1."):
                skipped_frontmatter = False
            else:
                i += 1
                continue
        if line.startswith("```"):
            previous_list_kind = None
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            previous_list_kind = None
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        if not line or line == "---":
            previous_list_kind = None
            i += 1
            continue
        if line.startswith("> **Operating instruction"):
            previous_list_kind = None
            clean = line[2:].replace("**Operating instruction for Codex:** ", "")
            add_callout(doc, clean)
            i += 1
            continue
        if line.startswith("> "):
            previous_list_kind = None
            add_callout(doc, line[2:].strip(), label="REFERENCE OBJECTIVE")
            i += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            previous_list_kind = None
            level = len(heading.group(1)) - 1
            title = heading.group(2)
            if title == "15. Implementation Phases and Work Packages":
                i += 1
                continue
            if level == 1 and (title.startswith("Phase ") or title.startswith("Appendix ")):
                doc.add_page_break()
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(p, title, default_size={1: 16, 2: 13, 3: 12}[level], default_color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
            for run in p.runs:
                run.bold = True
            paragraph_keep_with_next(p)
            i += 1
            continue
        bullet = re.match(r"^(\s*)-\s+(.*)$", line)
        numbered = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if bullet or numbered:
            match = bullet or numbered
            level = min(len(match.group(1)) // 2, 2)
            content = match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            if numbered and previous_list_kind != "numbered":
                active_decimal_id = new_numbering_instance(doc, decimal_num_id)
            apply_numbering(p, bullet_num_id if bullet else active_decimal_id, level)
            add_inline_markdown(p, content)
            previous_list_kind = "bullet" if bullet else "numbered"
            i += 1
            continue
        previous_list_kind = None
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        add_inline_markdown(p, line)
        i += 1


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)


def add_running_furniture(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("AGENTIC ML/DATA PLATFORM  |  IMPLEMENTATION PLAYBOOK")
        set_font(r, size=8.5, color=MUTED, bold=True)
        set_paragraph_border_bottom(p, color="D6DCE3", size="4", space="3")
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("Version 1.0  •  Page ")
        set_font(r, size=9, color=MUTED)
        add_page_field(p)


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    headings = []
    for line in text.splitlines():
        if re.match(r"^## (\d+\.|Appendix )", line):
            headings.append(line[3:])

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_id, decimal_id = configure_numbering(doc)
    add_cover(doc)
    add_contents(doc, headings)
    add_callout(
        doc,
        "Treat this document as the execution playbook. Follow phase gates in order, preserve every invariant, and update the progress ledger and ADRs as implementation proceeds.",
    )
    doc.add_paragraph()
    render_markdown(doc, text, bullet_id, decimal_id)
    configure_page(doc)
    add_running_furniture(doc)

    props = doc.core_properties
    props.title = "Agentic ML/Data Platform Implementation Playbook"
    props.subject = "Detailed implementation authority for the end-to-end runtime and harness platform"
    props.author = "Codex"
    props.keywords = "agentic platform, ML platform, runtime, harness, implementation, control plane"
    props.comments = "Generated from the canonical repository playbook."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
